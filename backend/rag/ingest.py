from pathlib import Path

import pymupdf
from bs4 import BeautifulSoup
from docx import Document


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".txt",
    ".md",
    ".markdown",
    ".docx",
    ".html",
    ".htm",
}


def extract_pdf(file_path: str) -> list[dict]:
    """
    Extract text from a PDF page by page.
    """

    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(
            f"File not found: {file_path}"
        )

    document = pymupdf.open(file_path)

    pages = []

    try:
        for page_number, page in enumerate(
            document,
            start=1,
        ):
            text = page.get_text("text").strip()

            if text:
                pages.append(
                    {
                        "document": path.name,
                        "page": page_number,
                        "text": text,
                    }
                )
    finally:
        document.close()

    return pages


def extract_docx(file_path: str) -> list[dict]:
    """
    Extract text from a DOCX document.

    DOCX does not have reliable page information,
    so the document is represented as page 1.
    """

    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(
            f"File not found: {file_path}"
        )

    document = Document(file_path)

    paragraphs = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            paragraphs.append(text)

    # Extract text from tables as well.
    for table in document.tables:
        for row in table.rows:
            cells = []

            for cell in row.cells:
                cell_text = cell.text.strip()

                if cell_text:
                    cells.append(cell_text)

            if cells:
                paragraphs.append(
                    " | ".join(cells)
                )

    text = "\n".join(paragraphs).strip()

    if not text:
        return []

    return [
        {
            "document": path.name,
            "page": 1,
            "text": text,
        }
    ]


def extract_text_file(file_path: str) -> list[dict]:
    """
    Extract text from TXT or Markdown files.
    """

    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(
            f"File not found: {file_path}"
        )

    text = path.read_text(
        encoding="utf-8",
        errors="replace",
    ).strip()

    if not text:
        return []

    return [
        {
            "document": path.name,
            "page": 1,
            "text": text,
        }
    ]


def extract_html(file_path: str) -> list[dict]:
    """
    Extract readable text from HTML.
    """

    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(
            f"File not found: {file_path}"
        )

    html = path.read_text(
        encoding="utf-8",
        errors="replace",
    )

    soup = BeautifulSoup(
        html,
        "html.parser",
    )

    # Remove non-content elements.
    for element in soup(
        [
            "script",
            "style",
            "noscript",
        ]
    ):
        element.decompose()

    text = soup.get_text(
        separator="\n",
        strip=True,
    )

    if not text:
        return []

    return [
        {
            "document": path.name,
            "page": 1,
            "text": text,
        }
    ]


def extract_document(file_path: str) -> list[dict]:
    """
    Automatically select the correct extractor.

    Supported:
        PDF
        DOCX
        TXT
        Markdown
        HTML
    """

    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(
            f"File not found: {file_path}"
        )

    extension = path.suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported document type: {extension}. "
            f"Supported formats: "
            f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if extension == ".pdf":
        return extract_pdf(file_path)

    if extension == ".docx":
        return extract_docx(file_path)

    if extension in {
        ".txt",
        ".md",
        ".markdown",
    }:
        return extract_text_file(file_path)

    if extension in {
        ".html",
        ".htm",
    }:
        return extract_html(file_path)

    raise ValueError(
        f"No extractor available for: {extension}"
    )